"""Order API router."""
from django.http import HttpResponse
from ninja import Router, File, UploadedFile, Form, Query

from apps.orders.schemas.input_schema import (
    CreateOrderSchema,
    UpdateOrderSchema,
    UpdateOrderStatusSchema,
    UpdateAssignedUsersSchema,
    UploadOrderImageSchema,
    OrderFilterSchema
)
from apps.orders.schemas.output_schema import (
    OrderOutSchema,
    OrderDetailSchema
)
from apps.orders.schemas.activity_schema import OrderActivitySchema
from apps.orders.services.service_a import OrderService
from core.responses.api_response import ApiResponse, ErrorResponse
from core.utils.pagination import PaginatedResponse
from core.authentication import JWTAuth

orders_router = Router(auth=JWTAuth())
order_service = OrderService()


@orders_router.get("/permissions", response={200: dict})
def get_user_permissions(request):
    """Get current user's permissions."""
    try:
        from core.enums.base_enum import UserRole

        user = request.auth

        allowed_statuses = UserRole.get_allowed_statuses(user.role)

        return 200, {
            "role": user.role,
            "role_label": UserRole.get_label(UserRole(user.role)),
            "allowed_statuses": [status.value for status in allowed_statuses],
            "can_create_order": user.role in [UserRole.ADMIN.value, UserRole.MANAGER.value, UserRole.SALE.value],
        }
    except Exception as e:
        return 200, {
            "role": "unknown",
            "role_label": "Unknown",
            "allowed_statuses": [],
            "can_create_order": False,
        }


@orders_router.post("/", response={201: OrderDetailSchema, 400: ErrorResponse})
def create_order(request, payload: CreateOrderSchema):
    """Create a new order."""
    try:
        user = request.auth

        order = order_service.create_order(payload, user)
        return 201, order
    except ValueError as e:
        return 400, {"detail": str(e)}
    except Exception as e:
        return 400, {"detail": f"Error creating order: {str(e)}"}


@orders_router.get("/", response=PaginatedResponse[OrderOutSchema])
def list_orders(
    request,
    filters: OrderFilterSchema = Query(...)
):
    """List all orders with filters."""
    try:
        user = request.auth

        orders, total = order_service.get_orders(
            filters,
            user_id=user.id if user and filters.assigned_to_me else None
        )

        return PaginatedResponse.create(
            items=orders,
            total=total,
            page=filters.page,
            page_size=filters.page_size
        )
    except Exception as e:
        print(f"Error in list_orders: {e}")
        import traceback
        traceback.print_exc()
        return PaginatedResponse.create(
            items=[],
            total=0,
            page=1,
            page_size=20
        )


@orders_router.get("/{order_id}", response={200: OrderDetailSchema, 404: ErrorResponse})
def get_order(request, order_id: int):
    """Get order by ID."""
    order = order_service.get_order_by_id(order_id)
    if not order:
        return 404, {"detail": f"Order with ID {order_id} not found"}
    return 200, order


@orders_router.patch("/{order_id}", response={200: OrderDetailSchema, 400: ErrorResponse, 403: ErrorResponse})
def update_order(request, order_id: int, payload: UpdateOrderSchema):
    """Update order details (items, customer info, fees, etc.)."""
    try:
        from core.enums.base_enum import UserRole

        user = request.auth

        # Only admin, manager, and sale can edit orders
        if user.role not in [UserRole.ADMIN.value, UserRole.MANAGER.value, UserRole.SALE.value]:
            return 403, {"detail": "Chỉ Admin, Manager và Sale mới có quyền sửa đơn hàng"}

        order = order_service.update_order(order_id, payload, user)
        return 200, order
    except ValueError as e:
        return 400, {"detail": str(e)}
    except Exception as e:
        return 400, {"detail": f"Error updating order: {str(e)}"}


@orders_router.patch("/{order_id}/status", response={200: OrderDetailSchema, 400: ErrorResponse, 403: ErrorResponse})
def update_order_status(request, order_id: int, payload: UpdateOrderStatusSchema):
    """Update order status."""
    try:
        from core.enums.base_enum import UserRole

        user = request.auth

        # Check permission before updating status
        order = order_service.get_order_by_id(order_id)
        if not order:
            return 400, {"detail": "Order not found"}

        # Check if user can transition from current status to new status
        if not UserRole.can_transition(user.role, order.status, payload.new_status):
            return 403, {"detail": f"Bạn không có quyền chuyển đơn từ '{order.status}' sang '{payload.new_status}'. Vai trò của bạn chỉ được phép làm các giai đoạn: {', '.join(UserRole.get_allowed_statuses(user.role))}"}

        order = order_service.update_order_status(order_id, payload, user)
        return 200, order
    except ValueError as e:
        return 400, {"detail": str(e)}
    except Exception as e:
        return 400, {"detail": f"Error updating status: {str(e)}"}


@orders_router.post("/{order_id}/images", response={201: dict, 400: ErrorResponse})
def upload_order_image(
    request,
    order_id: int,
    image: UploadedFile = File(...),
    payload: Form[UploadOrderImageSchema] = None
):
    """Upload image for order."""
    try:
        user = request.auth

        image_type = payload.image_type if payload else "other"

        order_image = order_service.upload_order_image(
            order_id=order_id,
            image_file=image,
            image_type=image_type,
            user=user
        )

        return 201, {
            "message": "Image uploaded successfully",
            "image_id": order_image.id,
            "image_url": order_image.image.url
        }
    except ValueError as e:
        return 400, {"detail": str(e)}


@orders_router.patch("/{order_id}/assigned-users", response={200: OrderDetailSchema, 400: ErrorResponse, 403: ErrorResponse})
def update_assigned_users(request, order_id: int, payload: UpdateAssignedUsersSchema):
    """Update assigned users for an order."""
    try:
        from core.enums.base_enum import UserRole

        user = request.auth

        # Only admin and manager can reassign
        if user.role not in [UserRole.ADMIN, UserRole.MANAGER]:
            return 403, {"detail": "Chỉ Admin và Manager mới có quyền phân công lại nhân viên"}

        order = order_service.update_assigned_users(order_id, payload.assigned_to_ids, user)
        return 200, order
    except ValueError as e:
        return 400, {"detail": str(e)}
    except Exception as e:
        return 400, {"detail": f"Error updating assigned users: {str(e)}"}


@orders_router.get("/statistics/summary", response=dict)
def get_order_statistics(request):
    """Get order statistics."""
    try:
        stats = order_service.get_order_statistics()
        return stats
    except Exception as e:
        return {"error": str(e)}


@orders_router.delete("/{order_id}/images/{image_id}", response={204: None, 404: ErrorResponse})
def delete_order_image(request, order_id: int, image_id: int):
    """Delete an order image."""
    try:
        from apps.orders.models import OrderImage
        from apps.orders.websocket_utils import broadcast_order_image_deleted
        from apps.orders.schemas.output_schema import OrderDetailSchema

        image = OrderImage.objects.filter(id=image_id, order_id=order_id).first()
        if not image:
            return 404, {"detail": f"Image with ID {image_id} not found in order {order_id}"}

        # Get order before deleting image
        order = order_service.get_order_by_id(order_id)

        # Delete the file from disk
        if image.image:
            image.image.delete(save=False)

        # Delete the database record
        image.delete()

        # Broadcast image deleted event with full order
        order_data = OrderDetailSchema.from_orm(order).model_dump(mode='json')
        broadcast_order_image_deleted(order_id, image_id, order_data)

        return 204, None
    except Exception as e:
        return 400, {"detail": str(e)}


@orders_router.get("/{order_id}/activities", response={200: list[OrderActivitySchema], 404: ErrorResponse})
def get_order_activities(request, order_id: int):
    """Get activity log for an order."""
    try:
        from apps.orders.models import OrderActivity

        order = order_service.get_order_by_id(order_id)
        if not order:
            return 404, {"detail": f"Order with ID {order_id} not found"}

        activities = OrderActivity.objects.filter(order=order).select_related('user').order_by('-created_at')
        return 200, list(activities)
    except Exception as e:
        return 400, {"detail": str(e)}


@orders_router.delete("/{order_id}", response={204: None, 404: ErrorResponse})
def delete_order(request, order_id: int):
    """Delete an order."""
    try:
        from apps.orders.websocket_utils import broadcast_order_deleted

        order = order_service.get_order_by_id(order_id)
        if not order:
            return 404, {"detail": f"Order with ID {order_id} not found"}

        order.delete()

        # Broadcast order deleted event
        broadcast_order_deleted(order_id)

        return 204, None
    except Exception as e:
        return 400, {"detail": str(e)}


@orders_router.get("/{order_id}/export-pdf", url_name="export_pdf")
def export_order_pdf(request, order_id: int, type: str = "order_bill", size: str = "K80"):
    """Export order as PDF."""
    from reportlab.lib.pagesizes import A4, A5
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from io import BytesIO
    import os
    from django.conf import settings

    try:
        order = order_service.get_order_by_id(order_id)
        if not order:
            return HttpResponse("Order not found", status=404)

        # Create PDF
        buffer = BytesIO()

        # Set page size
        if size == "A4":
            pagesize = A4
        elif size == "A5":
            pagesize = A5
        elif size == "K80":
            pagesize = (80*mm, 297*mm)
        else:  # K57
            pagesize = (57*mm, 297*mm)

        p = canvas.Canvas(buffer, pagesize=pagesize)
        width, height = pagesize

        # Use Helvetica (built-in font)
        p.setFont('Helvetica', 10)

        # Draw content
        y_position = height - 40

        # Header
        p.setFont('Helvetica-Bold', 14)
        title = {
            'order_bill': 'PHIEU DAT HANG',
            'weighing_receipt': 'PHIEU CAN HANG',
            'payment_bill': 'HOA DON THANH TOAN',
            'delivery_note': 'PHIEU GIAO HANG'
        }.get(type, 'HOA DON')

        p.drawCentredString(width/2, y_position, title)
        y_position -= 30

        # Order info
        p.setFont('Helvetica', 10)
        p.drawString(20, y_position, f"Ma don: {order.order_number}")
        y_position -= 20
        p.drawString(20, y_position, f"Khach hang: {order.customer.name}")
        y_position -= 20
        p.drawString(20, y_position, f"SDT: {order.customer.phone}")
        y_position -= 20
        p.drawString(20, y_position, f"Dia chi: {order.customer.address}")
        y_position -= 30

        # Items
        p.setFont('Helvetica-Bold', 10)
        p.drawString(20, y_position, "San pham:")
        y_position -= 20

        p.setFont('Helvetica', 9)
        for item in order.items.all():
            item_text = f"{item.product.name} - {item.quantity} {item.unit} x {item.price:,.0f}d = {item.total_price:,.0f}d"
            p.drawString(30, y_position, item_text)
            y_position -= 15

        y_position -= 10

        # Totals
        p.setFont('Helvetica', 10)
        p.drawString(20, y_position, f"Tam tinh: {order.subtotal:,.0f}d")
        y_position -= 15
        p.drawString(20, y_position, f"Phi van chuyen: {order.shipping_fee:,.0f}d")
        y_position -= 15
        p.drawString(20, y_position, f"Phi chip: {order.chip_fee:,.0f}d")
        y_position -= 20

        p.setFont('Helvetica-Bold', 12)
        p.drawString(20, y_position, f"Tong cong: {order.total:,.0f}d")

        p.showPage()
        p.save()

        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{order.order_number}_{type}.pdf"'
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return HttpResponse(f"Error: {str(e)}", status=500)


@orders_router.get("/{order_id}/export-word", url_name="export_word")
def export_order_word(request, order_id: int, type: str = "order_bill", size: str = "A4"):
    """Export order as Word document."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    try:
        order = order_service.get_order_by_id(order_id)
        if not order:
            return HttpResponse("Order not found", status=404)

        # Create Word document
        doc = Document()

        # Set page margins for size
        sections = doc.sections
        for section in sections:
            if size == "A5":
                section.page_height = Inches(8.27)
                section.page_width = Inches(5.83)
            # A4 is default
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Title
        title = {
            'order_bill': 'PHIẾU ĐẶT HÀNG',
            'weighing_receipt': 'PHIẾU CÂN HÀNG',
            'payment_bill': 'HÓA ĐƠN THANH TOÁN',
            'delivery_note': 'PHIẾU GIAO HÀNG'
        }.get(type, 'HÓA ĐƠN')

        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Order info
        doc.add_paragraph(f"Mã đơn: {order.order_number}")
        doc.add_paragraph(f"Khách hàng: {order.customer.name}")
        doc.add_paragraph(f"Số điện thoại: {order.customer.phone}")
        doc.add_paragraph(f"Địa chỉ: {order.customer.address}")

        if order.delivery_time:
            doc.add_paragraph(f"Thời gian giao: {order.delivery_time.strftime('%d/%m/%Y %H:%M')}")

        # Items table
        doc.add_paragraph()
        doc.add_heading('Sản phẩm:', level=2)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'STT'
        hdr_cells[1].text = 'Sản phẩm'
        hdr_cells[2].text = 'Số lượng'
        hdr_cells[3].text = 'Đơn giá'
        hdr_cells[4].text = 'Thành tiền'

        # Data rows
        for idx, item in enumerate(order.items.all(), 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = item.product.name
            row_cells[2].text = f"{item.quantity} {item.unit}"
            row_cells[3].text = f"{item.price:,.0f}đ"
            row_cells[4].text = f"{item.total_price:,.0f}đ"

        # Totals
        doc.add_paragraph()
        doc.add_paragraph(f"Tạm tính: {order.subtotal:,.0f}đ")
        doc.add_paragraph(f"Phí vận chuyển: {order.shipping_fee:,.0f}đ")
        doc.add_paragraph(f"Phí chip: {order.chip_fee:,.0f}đ")

        total_para = doc.add_paragraph()
        total_para.add_run(f"Tổng cộng: {order.total:,.0f}đ").bold = True

        # Notes
        if order.notes:
            doc.add_paragraph()
            doc.add_heading('Ghi chú:', level=2)
            doc.add_paragraph(order.notes)

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{order.order_number}_{type}.docx"'
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return HttpResponse(f"Error: {str(e)}", status=500)
